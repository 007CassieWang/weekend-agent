from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path("/Users/wangyaojie/Downloads/20238310570215-赵静-修改稿.docx")
OUT = Path("/Users/wangyaojie/Downloads/赵静-终稿.docx")

CN_NUM = "一二三四五六七八九十"
H1_RE = re.compile(rf"^[{CN_NUM}]+、")
H2_RE = re.compile(rf"^（[{CN_NUM}]+）")


def set_run_font(run, east_asia: str = "宋体", size_pt: float = 14, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_run(paragraph, text: str, font: str = "宋体", size: float = 14, bold: bool = False):
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold)
    return run


def set_paragraph(paragraph, align=None, first_indent=None, line_spacing=1.5, before=0, after=0) -> None:
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.first_line_indent = first_indent
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_text_with_breaks(paragraph, text: str, font: str, size: float, bold: bool = False) -> None:
    for idx, part in enumerate(text.split("\n")):
        if idx:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        if part:
            add_run(paragraph, part, font, size, bold)


def add_page_number_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    clear_runs(paragraph)
    set_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER, None, 1.0)

    run = paragraph.add_run()
    set_run_font(run, "宋体", 9)
    for kind, text in (("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)):
        if kind:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        else:
            el = OxmlElement("w:instrText" if text == " PAGE " else "w:t")
            if text == " PAGE ":
                el.set(qn("xml:space"), "preserve")
            el.text = text
        run._r.append(el)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    add_page_number_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def paragraph(doc: Document, text: str = "", font: str = "宋体", size: float = 14, bold: bool = False, **fmt):
    p = doc.add_paragraph()
    set_paragraph(p, **fmt)
    if text:
        add_text_with_breaks(p, text, font, size, bold)
    return p


def add_cover(doc: Document, title: str) -> None:
    paragraph(doc, "", line_spacing=1.0)
    paragraph(doc, "毕 业 作 业", "宋体", 31, True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, after=36)
    paragraph(doc, "", line_spacing=1.0)
    paragraph(doc, "毕业作业题目：", "宋体", 15, True, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5)
    paragraph(doc, title, "宋体", 16, True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=36)
    for label, value in [
        ("学院/分校：", ""),
        ("年级、专业：", "行政管理"),
        ("教育层次：", "专科"),
        ("学生姓名：", "赵静"),
        ("学    号：", "20238310570215"),
        ("指导教师：", ""),
        ("完成日期：", "年  月  日"),
    ]:
        p = paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Cm(3.0), line_spacing=1.8, after=3)
        add_run(p, label, "宋体", 14)
        add_run(p, value, "宋体", 14)
    doc.add_page_break()


def add_requirements_and_statement(doc: Document) -> None:
    paragraph(doc, "毕业作业基本要求", "宋体", 18, True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=12)
    reqs = [
        "1．毕业作业应在教师指导下，由学生个人独立完成。",
        "2. 对于以报告等形式进行的毕业作业选题，原则上应一人一题。对于综合性课题，确实需要有多人合作完成的话，则必须明确分工，由学生各自独立完成所分担的部分。严禁出现抄袭、代笔、剽窃等弄虚作假行为。",
    ]
    for req in reqs:
        paragraph(doc, req, "宋体", 14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=Cm(0.74), line_spacing=1.5, after=6)
    paragraph(doc, "*************************************", "宋体", 22, True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.2, before=18)
    paragraph(doc, "声  明", "宋体", 22, True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, before=12, after=12)
    statement = (
        "本人郑重声明：所呈交的毕业作业，是本人在老师指导下，独立进行研究所取得的成果。作业中除了已经注明引用的内容外，"
        "不包含任何他人享有的著作权内容。其他个人和集体对本研究工作的启发和所做出的贡献，均以明确的方式标明。"
        "如被查证有抄袭或剽窃行为，本人愿意承担由此引发的法律后果，并依据学校的规章制度接受相应处理。"
    )
    paragraph(doc, statement, "宋体", 14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=Cm(0.74), line_spacing=1.5)
    paragraph(doc, "签  名：                         日  期：      年   月   日", "宋体", 14, True, align=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.5, before=18)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    paragraph(doc, "目  录", "宋体", 14, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, before=24, after=24)
    entries = [
        ("摘  要", "I", 0),
        ("一、调查过程", "1", 0),
        ("（一）调查目的与意义", "1", 1),
        ("（二）调查对象概况", "1", 1),
        ("（三）调查时间", "1", 1),
        ("（四）调查方式", "1", 1),
        ("二、上海铵之包装材料有限公司员工培训工作的现状", "2", 0),
        ("（一）培训工作的责任机构与协同机制", "2", 1),
        ("（二）培训方式与内容体系", "2", 1),
        ("（三）培训工作的主要特点", "2", 1),
        ("三、调查的结论及思考", "3", 0),
        ("（一）员工培训工作的主要难点", "3", 1),
        ("（二）优化员工培训工作的对策思考", "3", 1),
    ]
    for title, page, level in entries:
        dots = "…" * max(4, 28 - len(title))
        paragraph(doc, f"{'  ' * level}{title}{dots}{page}", "宋体", 12, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5)
    doc.add_page_break()


def add_abstract(doc: Document, intro: str) -> None:
    paragraph(doc, "摘  要", "宋体", 14, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, before=24, after=24)
    abstract = intro.removeprefix("引言：")
    paragraph(doc, abstract, "宋体", 14, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=Cm(0.74), line_spacing=1.5)
    doc.add_page_break()


def add_body_paragraph(doc: Document, source_paragraph, idx: int) -> None:
    text = source_paragraph.text
    stripped = text.strip()
    p = doc.add_paragraph()
    if not stripped:
        set_paragraph(p, None, None, 1.5)
        return

    if idx == 0 or (idx == 4 and "调查报告" in stripped):
        set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, None, 1.5)
        add_text_with_breaks(p, text, "黑体", 16, True)
        return

    if idx >= 5 and H1_RE.match(stripped):
        set_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT, None, 1.5, before=12, after=12)
        add_text_with_breaks(p, text, "黑体", 18, True)
        return

    if H2_RE.match(stripped):
        set_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT, None, 1.5, before=9, after=0 if "\n" in text else 9)
        if "\n" in text:
            heading, rest = text.split("\n", 1)
            add_run(p, heading, "黑体", 15, True)
            p.add_run().add_break(WD_BREAK.LINE)
            add_text_with_breaks(p, rest, "宋体", 14, False)
        else:
            add_text_with_breaks(p, text, "黑体", 15, True)
        return

    set_paragraph(p, WD_ALIGN_PARAGRAPH.JUSTIFY, Cm(0.74), 1.5)
    add_text_with_breaks(p, text, "宋体", 14, False)


def build() -> None:
    source = Document(SRC)
    title = source.paragraphs[0].text.strip()
    intro = next((p.text for p in source.paragraphs if p.text.startswith("引言：")), "")

    doc = Document()
    configure_document(doc)
    add_cover(doc, title)
    add_requirements_and_statement(doc)
    add_contents(doc)
    add_abstract(doc, intro)
    for idx, src_p in enumerate(source.paragraphs):
        add_body_paragraph(doc, src_p, idx)
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
