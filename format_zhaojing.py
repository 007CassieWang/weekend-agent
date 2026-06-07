from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path("/Users/wangyaojie/Downloads/20238310570215-赵静-修改稿.docx")
OUT = Path("/Users/wangyaojie/Downloads/赵静-终稿.docx")

CN_NUM = "一二三四五六七八九十"
H1_RE = re.compile(rf"^[{CN_NUM}]+、")
H2_RE = re.compile(rf"^（[{CN_NUM}]+）")


def set_run_font(run, east_asia: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_spacing(paragraph, line_spacing: float = 1.5) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = line_spacing


def clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_text_with_breaks(paragraph, text: str, font: str, size_pt: float, bold: bool = False) -> None:
    parts = text.split("\n")
    for idx, part in enumerate(parts):
        if idx:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        if part:
            run = paragraph.add_run(part)
            set_run_font(run, font, size_pt, bold)


def add_page_number_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, 1.0)
    clear_runs(paragraph)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    for el in (fld_begin, instr, fld_sep, text, fld_end):
        run = paragraph.add_run()
        set_run_font(run, "宋体", 9)
        run._r.append(el)


def format_document() -> None:
    doc = Document(SRC)

    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)
        add_page_number_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        stripped = text.strip()
        set_paragraph_spacing(paragraph, 1.5)

        if not stripped:
            paragraph.paragraph_format.first_line_indent = None
            continue

        if idx == 0 or (idx == 4 and "调查报告" in stripped):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            clear_runs(paragraph)
            add_text_with_breaks(paragraph, text, "黑体", 16, True)
            continue

        if idx >= 5 and H1_RE.match(stripped):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)
            clear_runs(paragraph)
            add_text_with_breaks(paragraph, text, "黑体", 18, True)
            continue

        if H2_RE.match(stripped):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(9)
            paragraph.paragraph_format.space_after = Pt(9 if "\n" not in text else 0)
            clear_runs(paragraph)
            if "\n" in text:
                heading, rest = text.split("\n", 1)
                run = paragraph.add_run(heading)
                set_run_font(run, "黑体", 15, True)
                paragraph.add_run().add_break(WD_BREAK.LINE)
                add_text_with_breaks(paragraph, rest, "宋体", 14, False)
            else:
                add_text_with_breaks(paragraph, text, "黑体", 15, True)
            continue

        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(28)
        clear_runs(paragraph)
        add_text_with_breaks(paragraph, text, "宋体", 14, False)

    doc.save(OUT)


if __name__ == "__main__":
    format_document()
    print(OUT)
